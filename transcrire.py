#!/usr/bin/env python3
"""
Pipeline de transcription d'interviews YouTube
===============================================
Télécharge l'audio depuis YouTube, transcrit avec WhisperX + Pyannote,
corrige l'attribution des locuteurs (mapping intelligent + heuristiques +
passe Claude), puis envoie à Claude pour réécriture du français parlé
en français écrit élégant.

Par défaut, les interviewers sont Martin Bernard et Ouriel Barbezat
(Antithèse · Bon pour la tête). Le logo du média est intégré au docx.
Utilisez --interviewers / --hosts pour d'autres interviewers.

Usage:
    python transcrire.py URL --invites "Pierre Gallaz"
    python transcrire.py URL --invites "Sophie" "Paul" --skip 2:30
    python transcrire.py interview.mp4 --hosts "Jean-Marc" --invites "Sophie"
    python transcrire.py interview.mp3 --invites "Sophie" --output fichier.docx
    python transcrire.py --resume rewrite --invites "Sophie" --output fichier.docx

Pré-requis:
    - GPU NVIDIA avec CUDA (GTX 3090 recommandé)
    - Clé API Anthropic dans ANTHROPIC_API_KEY
    - Token Hugging Face dans HF_TOKEN (pour Pyannote)
    - Dépendances: pip install -r requirements.txt
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
import time
import urllib.request
from pathlib import Path

import hw
hw.setup_rocm_env()  # AMD/ROCm (gfx1151) : pose HSA_OVERRIDE_* avant tout import torch

# ---------------------------------------------------------------------------
# Défauts Antithèse · Bon pour la tête
# ---------------------------------------------------------------------------

DEFAULT_INTERVIEWERS = ["Martin Bernard", "Ouriel Barbezat"]
ANTITHESE_LOGO_URL = (
    "https://www.antithese.info/wp-content/uploads/Logo-Antihese-et-Bon-pour-la-tete.svg"
)
ANTITHESE_PLAYLIST_URL = (
    "https://www.youtube.com/playlist?list=PL8vbDh8Oio5WCWqqW61zqgroI4horNFsp"
)


# ---------------------------------------------------------------------------
# Playlist : récupération, affichage, sélection interactive
# ---------------------------------------------------------------------------

def fetch_playlist() -> list[dict]:
    """
    Récupère la liste des vidéos de la playlist Antithèse via yt-dlp.
    Retourne une liste de dicts : {url, raw_title, title, guests}
    """
    print("📋 Récupération de la playlist Antithèse...")
    cmd = [
        "yt-dlp", "--flat-playlist",
        "--print", "%(title)s\t%(url)s",
        ANTITHESE_PLAYLIST_URL,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"❌ Erreur yt-dlp : {result.stderr}")
        sys.exit(1)

    videos = []
    for line in result.stdout.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        parts = line.split("\t", 1)
        if len(parts) != 2:
            continue
        raw_title, url = parts

        if "|" in raw_title:
            title_part, guests_part = raw_title.split("|", 1)
            title = title_part.strip()
            # Split guests on ",", " et " or " & "
            guests_raw = re.split(r"\s*,\s*|\s+et\s+|\s*&\s*", guests_part.strip())
            guests = [g.strip() for g in guests_raw if g.strip()]
        else:
            # Certains titres utilisent " l " (L minuscule) comme séparateur
            # avant le nom de l'invité (majuscule), à distinguer de l' (article)
            sep_match = re.search(r" l (?=[A-ZÀ-Ü])", raw_title)
            if sep_match:
                title = raw_title[:sep_match.start()].strip()
                guests_part = raw_title[sep_match.end():].strip()
                guests_raw = re.split(r"\s*,\s*|\s+et\s+|\s*&\s*", guests_part)
                guests = [g.strip() for g in guests_raw if g.strip()]
            else:
                title = raw_title.strip()
                guests = ["Invité inconnu"]

        videos.append({
            "url": url,
            "raw_title": raw_title,
            "title": title,
            "guests": guests,
        })

    print(f"   ✅ {len(videos)} vidéo(s) trouvée(s)\n")
    return videos


def display_playlist(videos: list[dict]):
    """Affiche un tableau numéroté des vidéos de la playlist."""
    # Calculer les largeurs de colonnes
    col_guests = max(
        (len(", ".join(v["guests"])) for v in videos), default=10,
    )
    col_guests = max(col_guests, 10)
    col_title = max((len(v["title"]) for v in videos), default=10)
    col_title = max(col_title, 10)

    # En-tête
    print(f"  {'N°':>4}   {'Invité(s)':<{col_guests}}   {'Titre':<{col_title}}")
    print(f"  {'───':>4}   {'─' * col_guests}   {'─' * col_title}")

    for i, v in enumerate(videos, 1):
        guests_str = ", ".join(v["guests"])
        print(f"  {i:>4}   {guests_str:<{col_guests}}   {v['title']:<{col_title}}")

    print()


def select_videos(videos: list[dict]) -> list[dict]:
    """
    Prompt interactif pour sélectionner des vidéos.
    Accepte : all, 1-5 (range inclusif), 1, 3, 5 (indices séparés par virgule).
    """
    while True:
        try:
            choice = input("Sélection (all / 1-5 / 1, 3, 5) : ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\n❌ Sélection annulée.")
            sys.exit(0)

        if not choice:
            continue

        if choice == "all":
            return list(videos)

        # Range : "1-5"
        range_match = re.match(r"^(\d+)\s*-\s*(\d+)$", choice)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            if start < 1 or end > len(videos) or start > end:
                print(f"   ⚠️  Plage invalide (1-{len(videos)} attendu)")
                continue
            return [videos[i] for i in range(start - 1, end)]

        # Comma-separated : "1, 3, 5"
        try:
            indices = [int(x.strip()) for x in choice.split(",")]
        except ValueError:
            print("   ⚠️  Format invalide. Exemples : all, 1-5, 1, 3, 5")
            continue

        invalid = [i for i in indices if i < 1 or i > len(videos)]
        if invalid:
            print(f"   ⚠️  Indice(s) hors limites : {invalid} (1-{len(videos)} attendu)")
            continue

        return [videos[i - 1] for i in indices]


# ---------------------------------------------------------------------------
# Utilitaires : chronomètre, compteur API
# ---------------------------------------------------------------------------

class StepTimer:
    """Chronomètre pour mesurer la durée de chaque étape du pipeline."""

    def __init__(self):
        self.steps = []
        self.current_step = None
        self.current_start = None
        self.total_start = time.time()

    def start(self, name: str):
        if self.current_step:
            self.stop()
        self.current_step = name
        self.current_start = time.time()

    def stop(self):
        if self.current_step and self.current_start:
            elapsed = time.time() - self.current_start
            self.steps.append((self.current_step, elapsed))
            self.current_step = None
            self.current_start = None

    @staticmethod
    def format_duration(seconds: float) -> str:
        m, s = divmod(int(seconds), 60)
        return f"{m}:{s:02d}"

    def summary(self) -> str:
        self.stop()
        lines = []
        for name, elapsed in self.steps:
            lines.append(f"   {name:40s} {self.format_duration(elapsed)}")
        total = time.time() - self.total_start
        lines.append(f"   {'TOTAL':40s} {self.format_duration(total)}")
        return "\n".join(lines)


class APIStats:
    """Compteur d'appels et de tokens API."""

    def __init__(self):
        self.calls = 0
        self.input_tokens = 0
        self.output_tokens = 0

    def record(self, message):
        self.calls += 1
        if hasattr(message, "usage"):
            self.input_tokens += getattr(message.usage, "input_tokens", 0)
            self.output_tokens += getattr(message.usage, "output_tokens", 0)

    def estimated_cost(self) -> float:
        # Sonnet 4 pricing approx: $3/MTok in, $15/MTok out
        return (self.input_tokens * 3 + self.output_tokens * 15) / 1_000_000

    def summary(self) -> str:
        cost = self.estimated_cost()
        return (
            f"   Appels API          : {self.calls}\n"
            f"   Tokens entrée       : {self.input_tokens:,}\n"
            f"   Tokens sortie       : {self.output_tokens:,}\n"
            f"   Coût estimé         : ${cost:.2f}"
        )


# ---------------------------------------------------------------------------
# Checkpoints : sauvegarde / reprise d'état
# ---------------------------------------------------------------------------

def get_workdir(output_path: str) -> str:
    """Crée et retourne le répertoire de travail dédié."""
    base = Path(output_path).stem
    workdir = Path(output_path).parent / f"{base}_workdir"
    workdir.mkdir(parents=True, exist_ok=True)
    return str(workdir)


def save_checkpoint(workdir: str, step: str, data, step_number: int):
    """Sauvegarde un checkpoint JSON numéroté dans le workdir."""
    filename = f"{step_number:02d}_{step}.json"
    path = os.path.join(workdir, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"   💾 Checkpoint : {filename}")


def load_checkpoint(workdir: str, step: str, step_number: int):
    """Charge un checkpoint JSON depuis le workdir. Retourne None si absent."""
    filename = f"{step_number:02d}_{step}.json"
    path = os.path.join(workdir, filename)
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Statistiques de parole par locuteur
# ---------------------------------------------------------------------------

def compute_speaker_stats(segments: list[dict]) -> dict:
    """
    Calcule par locuteur : temps total, mots, questions, ratio, longueur
    moyenne des segments. Utilise les timestamps start/end des segments.
    """
    stats = {}

    for seg in segments:
        speaker = seg["speaker"]
        text = seg.get("text", "").strip()
        start = seg.get("start", 0)
        end = seg.get("end", 0)

        if speaker not in stats:
            stats[speaker] = {
                "total_time": 0.0,
                "word_count": 0,
                "num_questions": 0,
                "num_segments": 0,
            }

        words = len(text.split())
        is_q = _is_question_simple(text)

        stats[speaker]["total_time"] += max(end - start, 0)
        stats[speaker]["word_count"] += words
        stats[speaker]["num_questions"] += 1 if is_q else 0
        stats[speaker]["num_segments"] += 1

    for s in stats.values():
        s["question_ratio"] = s["num_questions"] / max(s["num_segments"], 1)
        s["avg_segment_length"] = s["word_count"] / max(s["num_segments"], 1)

    return stats


def _is_question_simple(text: str) -> bool:
    """Détection simple de question (pour le calcul de stats)."""
    text = text.strip()
    if text.endswith("?"):
        return True
    starters = [
        "comment", "pourquoi", "qu'est-ce", "quel", "quelle", "quels", "quelles",
        "est-ce que", "est-ce qu'", "combien", "où", "quand", "qui est",
        "que pensez", "que penses", "qu'en pensez", "qu'en penses",
        "pouvez-vous", "peux-tu", "pourriez-vous",
        "parlez-nous", "expliquez", "diriez-vous", "racontez",
    ]
    lower = text.lower()
    return any(lower.startswith(q) for q in starters)


def format_speaker_stats(stats: dict) -> str:
    """Formate les statistiques pour affichage ou inclusion dans un prompt."""
    lines = []
    for speaker, s in sorted(stats.items(), key=lambda x: x[1]["total_time"], reverse=True):
        time_min = s["total_time"] / 60
        lines.append(
            f"  {speaker}: {time_min:.1f}min, {s['word_count']} mots, "
            f"{s['num_segments']} segments (moy. {s['avg_segment_length']:.0f} mots), "
            f"{s['num_questions']} questions ({s['question_ratio']:.0%})"
        )
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 1. Téléchargement audio depuis YouTube
# ---------------------------------------------------------------------------

def download_audio(url_or_path: str, output_dir: str) -> str:
    """Télécharge l'audio depuis YouTube ou prépare un fichier local."""

    if os.path.isfile(url_or_path):
        print(f"📁 Fichier local détecté : {url_or_path}")
        ext = Path(url_or_path).suffix.lower()

        if ext == ".wav":
            return url_or_path

        print(f"🔄 Conversion {ext} → WAV...")
        output_path = os.path.join(output_dir, "audio.wav")
        cmd = [
            "ffmpeg", "-i", url_or_path,
            "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", "-y",
            output_path,
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"❌ Erreur ffmpeg : {result.stderr}")
            sys.exit(1)

        print(f"✅ Audio converti : {output_path}")
        return output_path

    print("⬇️  Téléchargement audio depuis YouTube...")
    output_path = os.path.join(output_dir, "audio.wav")

    cmd = [
        "yt-dlp",
        "--extract-audio",
        "--audio-format", "wav",
        "--audio-quality", "0",
        "--output", os.path.join(output_dir, "audio.%(ext)s"),
        "--no-playlist",
        url_or_path,
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"❌ Erreur yt-dlp : {result.stderr}")
        sys.exit(1)

    for f in os.listdir(output_dir):
        if f.startswith("audio") and f.endswith(".wav"):
            output_path = os.path.join(output_dir, f)
            break

    print(f"✅ Audio téléchargé : {output_path}")
    return output_path


def trim_audio(audio_path: str, ignore_str: str, output_dir: str) -> str:
    """Coupe le début de l'audio à partir de MM:SS."""
    parts = ignore_str.split(":")
    if len(parts) == 2:
        minutes, seconds = int(parts[0]), int(parts[1])
    elif len(parts) == 1:
        minutes, seconds = 0, int(parts[0])
    else:
        print(f"❌ Format invalide : {ignore_str} (attendu MM:SS)")
        sys.exit(1)

    total_seconds = minutes * 60 + seconds
    start_time = f"{minutes:02d}:{seconds:02d}"

    print(f"✂️  Suppression des {start_time} premières minutes...")

    trimmed_path = os.path.join(output_dir, "audio_trimmed.wav")
    cmd = [
        "ffmpeg", "-i", audio_path,
        "-ss", str(total_seconds),
        "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", "-y",
        trimmed_path,
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"❌ Erreur ffmpeg : {result.stderr}")
        sys.exit(1)

    print(f"✅ Audio tronqué (début à {start_time})")
    return trimmed_path


# ---------------------------------------------------------------------------
# 2. Transcription + Diarisation avec WhisperX
# ---------------------------------------------------------------------------

def transcribe_and_diarize(audio_path: str, num_speakers: int, hf_token: str,
                           whisper_model: str = "large-v3",
                           batch_size: int = 16) -> list[dict]:
    """
    Transcrit l'audio avec WhisperX et diarise avec Pyannote.
    Retourne une liste de segments avec timestamps :
    [{"speaker": "SPEAKER_00", "text": "...", "start": 0.0, "end": 1.5}]
    """
    import whisperx
    import torch

    device = hw.device()  # « cuda » couvre CUDA et ROCm/HIP
    compute_type = hw.whisper_compute_type()

    print(f"🎙️  Transcription avec WhisperX (device={device}, modèle={whisper_model}, batch={batch_size})...")

    # Étape 1 : Transcription
    model = whisperx.load_model(whisper_model, device, compute_type=compute_type, language="fr")
    audio = whisperx.load_audio(audio_path)
    result = model.transcribe(audio, batch_size=batch_size, language="fr")

    import gc
    del model
    gc.collect()
    if device == "cuda":
        torch.cuda.empty_cache()

    print("📐 Alignement des mots...")

    # Étape 2 : Alignement
    align_model, metadata = whisperx.load_align_model(language_code="fr", device=device)
    result = whisperx.align(
        result["segments"], align_model, metadata, audio, device,
        return_char_alignments=False,
    )

    del align_model
    gc.collect()
    if device == "cuda":
        torch.cuda.empty_cache()

    print(f"👥 Diarisation ({num_speakers} locuteurs)...")

    # Étape 3 : Diarisation
    from whisperx.diarize import DiarizationPipeline
    diarize_model = DiarizationPipeline(token=hf_token, device=device)
    diarize_segments = diarize_model(
        audio, min_speakers=num_speakers, max_speakers=num_speakers,
    )

    # Étape 4 : Attribution des locuteurs aux segments
    result = whisperx.assign_word_speakers(diarize_segments, result)

    del diarize_model
    gc.collect()
    if device == "cuda":
        torch.cuda.empty_cache()

    # Consolider les segments consécutifs du même locuteur (en gardant timestamps)
    segments = result["segments"]
    consolidated = []

    for seg in segments:
        speaker = seg.get("speaker", "UNKNOWN")
        text = seg.get("text", "").strip()
        start = seg.get("start", 0.0)
        end = seg.get("end", 0.0)
        if not text:
            continue

        if consolidated and consolidated[-1]["speaker"] == speaker:
            consolidated[-1]["text"] += " " + text
            consolidated[-1]["end"] = end
        else:
            consolidated.append({
                "speaker": speaker,
                "text": text,
                "start": start,
                "end": end,
            })

    print(f"✅ Transcription terminée : {len(consolidated)} blocs de parole")
    return consolidated


# ---------------------------------------------------------------------------
# 3. Mapping intelligent des locuteurs
# ---------------------------------------------------------------------------

def replace_speaker_names(segments: list[dict], interviewers: list[str],
                          invites: list[str]) -> list[dict]:
    """
    Remplace SPEAKER_00, SPEAKER_01, etc. par les vrais noms.
    Utilise l'analyse des patterns de parole (temps, questions, longueur)
    au lieu du simple ordre d'apparition.
    """
    all_names = interviewers + invites

    # Identifier les IDs de locuteurs uniques (ordre d'apparition)
    speaker_ids = []
    for seg in segments:
        spk = seg["speaker"]
        if spk not in speaker_ids:
            speaker_ids.append(spk)

    num_detected = len(speaker_ids)
    num_expected = len(all_names)

    if num_detected != num_expected:
        print(f"   ⚠️  Attention : {num_detected} locuteurs détectés mais {num_expected} attendus")

    # Statistiques par identifiant brut
    stats = compute_speaker_stats(segments)

    # Score « invité » : temps long + segments longs − ratio questions élevé
    # Plus le score est élevé, plus le locuteur est probablement un invité
    scores = {}
    for spk_id in speaker_ids:
        s = stats.get(spk_id)
        if not s:
            scores[spk_id] = 0
            continue
        scores[spk_id] = (
            s["avg_segment_length"] * 0.4
            + s["total_time"] * 0.3
            - s["question_ratio"] * 100 * 0.3
        )

    # Tri : scores les plus bas → interviewers, les plus hauts → invités
    sorted_speakers = sorted(speaker_ids, key=lambda x: scores.get(x, 0))

    interviewer_ids = sorted_speakers[:len(interviewers)]
    invite_ids = sorted_speakers[len(interviewers):]

    # Parmi les interviewers, celui avec le plus haut ratio de questions = principal
    interviewer_ids.sort(
        key=lambda x: stats.get(x, {}).get("question_ratio", 0), reverse=True,
    )
    # Parmi les invités, celui avec le plus de temps de parole = premier listé
    invite_ids.sort(
        key=lambda x: stats.get(x, {}).get("total_time", 0), reverse=True,
    )

    mapping = {}
    for i, spk_id in enumerate(interviewer_ids):
        mapping[spk_id] = interviewers[i] if i < len(interviewers) else f"Interviewer {i+1}"
    for i, spk_id in enumerate(invite_ids):
        mapping[spk_id] = invites[i] if i < len(invites) else f"Invité {i+1}"

    # Locuteurs en surplus
    for spk_id in speaker_ids:
        if spk_id not in mapping:
            mapping[spk_id] = f"Locuteur ({spk_id})"

    # Affichage du raisonnement détaillé
    print("🏷️  Attribution intelligente des noms :")
    for spk_id in speaker_ids:
        s = stats.get(spk_id, {})
        time_min = s.get("total_time", 0) / 60
        role = "interviewer" if spk_id in interviewer_ids else "invité"
        print(f"   {spk_id} → {mapping[spk_id]} ({role})")
        print(
            f"      temps={time_min:.1f}min, mots={s.get('word_count', 0)}, "
            f"segments={s.get('num_segments', 0)} (moy. {s.get('avg_segment_length', 0):.0f} mots), "
            f"questions={s.get('num_questions', 0)} ({s.get('question_ratio', 0):.0%}), "
            f"score invité={scores.get(spk_id, 0):.1f}"
        )

    for seg in segments:
        seg["speaker"] = mapping.get(seg["speaker"], seg["speaker"])

    return segments


# ---------------------------------------------------------------------------
# 4. Détection de questions et relances (helpers pour heuristiques)
# ---------------------------------------------------------------------------

def is_question(text: str) -> bool:
    """Détecte si un segment est probablement une question."""
    text = text.strip()
    if text.endswith("?"):
        return True
    question_starters = [
        "comment", "pourquoi", "qu'est-ce", "quel", "quelle", "quels", "quelles",
        "est-ce que", "est-ce qu'", "combien", "où", "quand", "qui est",
        "que pensez", "que penses", "qu'en pensez", "qu'en penses",
        "pouvez-vous", "peux-tu", "pourriez-vous",
        "parlez-nous", "expliquez", "diriez-vous", "racontez",
        "et alors", "et donc", "c'est-à-dire",
        "décrivez", "comment ça", "dans quelle mesure",
        "à quel point", "de quelle manière",
    ]
    lower = text.lower()
    return any(lower.startswith(q) for q in question_starters)


def is_relance(text: str) -> bool:
    """Détecte les relances courtes typiques d'un interviewer."""
    stripped = text.strip().rstrip(".!?…").lower()
    relances = [
        "d'accord", "ok", "oui", "bien sûr", "tout à fait",
        "intéressant", "je vois", "très bien", "effectivement",
        "absolument", "exactement", "en effet", "ah oui",
        "ah bon", "vraiment", "c'est vrai", "hmm",
        "et ensuite", "et après", "et donc",
        "je comprends", "fascinant", "formidable",
    ]
    if stripped in relances:
        return True
    for r in relances:
        if stripped.startswith(r) and len(stripped.split()) <= 8:
            return True
    return False


def find_nearest_interviewer(segments: list[dict], index: int,
                             interviewers: list[str]) -> str:
    """Trouve l'interviewer le plus proche dans le contexte."""
    for offset in range(1, min(10, len(segments))):
        for direction in (-1, 1):
            j = index + direction * offset
            if 0 <= j < len(segments) and segments[j]["speaker"] in interviewers:
                return segments[j]["speaker"]
    return interviewers[0]


# ---------------------------------------------------------------------------
# 5. Correction heuristique des attributions
# ---------------------------------------------------------------------------

def fix_attributions_heuristic(segments: list[dict], interviewers: list[str],
                               invites: list[str]) -> list[dict]:
    """
    Corrige les cas évidents de mauvaise attribution basés sur la structure
    d'une interview : questions courtes = interviewers, réponses longues = invités.
    """

    def wc(text: str) -> int:
        return len(text.split())

    corrections = 0

    for i, seg in enumerate(segments):
        speaker = seg["speaker"]
        text = seg["text"]
        words = wc(text)

        # Règle 1 : Relance courte (< 8 mots) attribuée à un invité → interviewer
        if speaker in invites and is_relance(text) and words < 8:
            seg["speaker"] = find_nearest_interviewer(segments, i, interviewers)
            corrections += 1
            continue

        # Règle 2 : Question courte (< 40 mots) attribuée à un invité,
        #           suivie d'une longue réponse → interviewer
        if speaker in invites and is_question(text) and words < 40:
            next_is_long = (
                i + 1 < len(segments)
                and wc(segments[i + 1]["text"]) > 50
            )
            if next_is_long:
                seg["speaker"] = find_nearest_interviewer(segments, i, interviewers)
                corrections += 1
                continue

        # Règle 3 : Question très courte (< 15 mots) attribuée à un invité
        if speaker in invites and is_question(text) and words < 15:
            seg["speaker"] = find_nearest_interviewer(segments, i, interviewers)
            corrections += 1
            continue

        # Règle 4 : Segment très long (> 200 mots) attribué à un interviewer
        #           → probablement un invité (sauf introductions, i > 2)
        if speaker in interviewers and words > 200 and i > 2:
            for offset in range(1, min(10, len(segments))):
                found = False
                for direction in (-1, 1):
                    j = i + direction * offset
                    if 0 <= j < len(segments) and segments[j]["speaker"] in invites:
                        seg["speaker"] = segments[j]["speaker"]
                        corrections += 1
                        found = True
                        break
                if found:
                    break
            continue

        # Règle 5 : Deux segments consécutifs du même invité,
        #           le premier est court (< 30 mots) suivi d'un long (> 80 mots)
        #           → le court est une relance d'interviewer
        if (i + 1 < len(segments)
                and speaker in invites
                and segments[i + 1]["speaker"] == speaker
                and words < 30 and wc(segments[i + 1]["text"]) > 80):
            if is_question(text) or is_relance(text):
                seg["speaker"] = find_nearest_interviewer(segments, i, interviewers)
                corrections += 1
                continue

        # Règle 6 : Deux segments consécutifs du même invité,
        #           le premier est une question (< 60 mots) → interviewer
        if (i + 1 < len(segments)
                and speaker in invites
                and segments[i + 1]["speaker"] == speaker
                and is_question(text) and words < 60):
            seg["speaker"] = find_nearest_interviewer(segments, i, interviewers)
            corrections += 1

    # Re-consolider après corrections (fusionner segments consécutifs même locuteur)
    consolidated = []
    for seg in segments:
        if consolidated and consolidated[-1]["speaker"] == seg["speaker"]:
            consolidated[-1]["text"] += " " + seg["text"]
            consolidated[-1]["end"] = seg.get("end", consolidated[-1].get("end", 0))
        else:
            consolidated.append(dict(seg))

    if corrections > 0:
        print(f"🔧 Corrections heuristiques : {corrections} attribution(s) corrigée(s)")
    else:
        print("🔧 Corrections heuristiques : aucune correction nécessaire")

    return consolidated


# ---------------------------------------------------------------------------
# 6. Chunking par tours de parole
# ---------------------------------------------------------------------------

def chunk_transcript(text: str, max_words: int = 3000,
                     context_turns: int = 2) -> list[str]:
    """
    Découpe la transcription en chunks en respectant les frontières de locuteur.
    Inclut les derniers tours du chunk précédent marqués [CONTEXTE PRÉCÉDENT].
    """
    blocks = text.split("\n\n")

    if len(text.split()) <= max_words:
        return [text]

    chunks = []
    current_blocks = []
    current_words = 0

    for block in blocks:
        block_words = len(block.split())

        if current_words + block_words > max_words and current_blocks:
            chunks.append("\n\n".join(current_blocks))

            # Chevauchement contextuel : derniers N tours
            context = current_blocks[-context_turns:]
            current_blocks = []
            current_words = 0

            if context:
                ctx_text = "\n\n".join(context)
                current_blocks.append(
                    f"[CONTEXTE PRÉCÉDENT]\n{ctx_text}\n[FIN CONTEXTE]"
                )
                current_words = sum(len(b.split()) for b in context)

        current_blocks.append(block)
        current_words += block_words

    if current_blocks:
        chunks.append("\n\n".join(current_blocks))

    return chunks


# ---------------------------------------------------------------------------
# 7. Appel API Claude avec retry et backoff exponentiel
# ---------------------------------------------------------------------------

def call_claude_api(client, model: str, system: str, user_msg: str,
                    api_stats: "APIStats", max_tokens: int = 8192,
                    max_retries: int = 3) -> str:
    """Appel API Claude avec retry, backoff exponentiel et validation."""
    delays = [2, 4, 8]

    for attempt in range(max_retries):
        try:
            message = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system,
                messages=[{"role": "user", "content": user_msg}],
            )

            api_stats.record(message)

            result = message.content[0].text
            if not result or not result.strip():
                raise ValueError("Réponse API vide")

            return result

        except Exception as e:
            if attempt < max_retries - 1:
                delay = delays[attempt]
                print(f"   ⚠️  Erreur API (tentative {attempt+1}/{max_retries}): {e}")
                print(f"   ⏳ Nouvelle tentative dans {delay}s...")
                time.sleep(delay)
            else:
                print(f"   ❌ Échec après {max_retries} tentatives: {e}")
                raise


# ---------------------------------------------------------------------------
# 8. Correction des attributions par Claude (passe dédiée)
# ---------------------------------------------------------------------------

def fix_attributions_claude(transcript: str, interviewers: list[str],
                            invites: list[str], api_key: str, model: str,
                            api_stats: "APIStats",
                            speaker_stats_str: str = "") -> str:
    """
    Passe dédiée : Claude corrige les attributions sans réécrire le texte.
    """
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    interviewers_str = ", ".join(interviewers)
    invites_str = ", ".join(invites)

    system_prompt = f"""Tu es un correcteur spécialisé dans l'attribution des prises de parole \
dans les transcriptions d'interviews.

CONTEXTE :
- Les INTERVIEWERS sont : {interviewers_str}
  {interviewers[0]} est l'interviewer principal qui pose la grande majorité des questions.
  Les autres interviewers interviennent ponctuellement.
- Les INVITÉS sont : {invites_str}
  Ils répondent aux questions. Ils font parfois de longues réponses de plusieurs paragraphes.

STATISTIQUES DE PAROLE :
{speaker_stats_str}

RÈGLES pour corriger les attributions :
1. Les QUESTIONS sont presque toujours posées par un interviewer.
   Exception : un invité peut poser une question rhétorique au milieu de sa réponse
   (ex: "Pourquoi ? Parce que..." ou "Qu'est-ce que ça veut dire ? Ça veut dire que...").
2. Les RÉPONSES longues et développées sont celles des invités.
3. Les relances courtes ("D'accord", "Intéressant", "Et ensuite ?", "Je vois") viennent \
des interviewers.
4. Un invité ne pose presque jamais de questions à l'interviewer.
5. Si une question isolée est attribuée à un invité (et n'est pas une question rhétorique \
suivie immédiatement de sa propre réponse), corrige-la vers l'interviewer le plus proche.
6. Si deux prises de parole consécutives sont attribuées au même invité et que la première \
est une question suivie d'une longue réponse, sépare-les : la question est de l'interviewer.

EXEMPLES :
  ❌ INCORRECT : "{invites[0]} : Comment en êtes-vous arrivé là ?"
     suivi de  : "{invites[0]} : J'ai commencé par..."
  ✅ CORRECT   : "{interviewers[0]} : Comment en êtes-vous arrivé là ?"
     suivi de  : "{invites[0]} : J'ai commencé par..."

  ❌ INCORRECT : "{invites[0]} : D'accord, et ensuite ?"
  ✅ CORRECT   : "{interviewers[0]} : D'accord, et ensuite ?"

  ✅ CORRECT (question rhétorique, NE PAS corriger) :
     "{invites[0]} : ...et c'est là que je me suis demandé : pourquoi est-ce qu'on fait ça ? \
En fait, la raison est simple..."

INSTRUCTIONS :
- Retourne la transcription COMPLÈTE avec les attributions corrigées.
- Ne modifie PAS le texte lui-même, UNIQUEMENT les noms des locuteurs.
- Conserve le format exact : "Nom : texte"
- Si le texte contient des sections [CONTEXTE PRÉCÉDENT]...[FIN CONTEXTE], \
ne les inclus PAS dans ta sortie. Elles sont là uniquement pour te donner du contexte.
- Si aucune correction n'est nécessaire, retourne le texte tel quel."""

    chunks = chunk_transcript(transcript)
    print(f"🔍 Correction des attributions par Claude ({len(chunks)} partie(s))...")

    corrected_parts = []
    for i, chunk in enumerate(chunks):
        if len(chunks) > 1:
            print(f"   Partie {i+1}/{len(chunks)}...")

        user_msg = f"""Voici une transcription d'interview avec des attributions potentiellement erronées.
Corrige UNIQUEMENT les attributions de locuteurs, sans modifier le texte.

---
{chunk}
---"""

        result = call_claude_api(client, model, system_prompt, user_msg, api_stats)
        corrected_parts.append(result)

        if i < len(chunks) - 1:
            time.sleep(1)

    result = "\n\n".join(corrected_parts)
    print("✅ Correction des attributions terminée")
    return result


# ---------------------------------------------------------------------------
# 9. Réécriture avec Claude API
# ---------------------------------------------------------------------------

def format_transcript_for_claude(segments: list[dict]) -> str:
    """Formate la transcription brute pour l'envoi à Claude."""
    lines = []
    for seg in segments:
        lines.append(f"{seg['speaker']} : {seg['text']}")
    return "\n\n".join(lines)


def rewrite_with_claude(transcript: str, interviewers: list[str],
                        invites: list[str], api_key: str, model: str,
                        api_stats: "APIStats") -> str:
    """Envoie la transcription à Claude pour réécriture parlé → écrit."""
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)

    interviewers_str = ", ".join(interviewers)
    invites_str = ", ".join(invites)

    system_prompt = f"""Tu es un éditeur professionnel spécialisé dans la transformation \
d'interviews orales en textes écrits élégants en français.

CONTEXTE :
- Interviewers : {interviewers_str}
- Invité(s) : {invites_str}

RÈGLES STRICTES :
1. Transforme le français parlé en français écrit fluide et naturel.
2. Supprime les hésitations, répétitions, faux départs, "euh", "ben", "du coup", "en fait" \
(quand ils sont des tics de langage).
3. Reformule les phrases bancales en phrases bien construites et agréables à lire.
4. PRÉSERVE FIDÈLEMENT le sens, les idées, le ton et la personnalité de chaque intervenant.
5. Ne résume JAMAIS. Ne coupe AUCUN contenu substantiel. Chaque idée exprimée doit rester.
6. Garde les expressions colorées ou personnelles qui font le charme de l'oral.
7. Maintiens la structure question/réponse avec les noms des intervenants.
8. Ne rajoute RIEN qui n'a pas été dit.
9. Pour les auto-corrections ("enfin, je veux dire...", "non en fait..."), garde la version \
corrigée par le locuteur et supprime le faux départ.
10. Préserve exactement : citations, noms propres, dates, chiffres, acronymes.
11. Pour les incises orales ("comment dire", "si je puis dire", "disons"), supprime-les \
si elles n'apportent rien, garde-les si elles font partie du style du locuteur.
12. Ne fusionne JAMAIS les prises de parole de deux locuteurs différents.

FORMAT DE SORTIE :
- Chaque prise de parole commence par le prénom complet en gras : **Nom**
- Un saut de ligne entre chaque prise de parole.
- Pas de timestamps.
- Pas de notes éditoriales.
- Si le texte contient des sections [CONTEXTE PRÉCÉDENT]...[FIN CONTEXTE], \
ne les inclus PAS dans ta sortie. Elles sont là uniquement pour te donner du contexte."""

    chunks = chunk_transcript(transcript)
    print(f"✍️  Réécriture avec Claude ({len(chunks)} partie(s))...")

    rewritten_parts = []
    for i, chunk in enumerate(chunks):
        if len(chunks) > 1:
            print(f"   Partie {i+1}/{len(chunks)}...")

        user_msg = f"""Voici la transcription de l'interview.
Réécris-la en suivant les règles indiquées.

---
{chunk}
---"""

        result = call_claude_api(client, model, system_prompt, user_msg, api_stats)
        rewritten_parts.append(result)

        if i < len(chunks) - 1:
            time.sleep(1)

    result = "\n\n".join(rewritten_parts)
    print("✅ Réécriture terminée")
    return result


# ---------------------------------------------------------------------------
# 10. Génération du document .docx
# ---------------------------------------------------------------------------

def generate_docx(rewritten_text: str, interviewers: list[str], invites: list[str],
                  output_path: str, video_url: str = None,
                  logo_path: str = None, is_antithese: bool = False,
                  titre: str = None):
    """Génère un fichier .docx élégant à partir du texte réécrit."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    # ── Palette de couleurs ──────────────────────────────────────────────
    COLOR_TITLE       = RGBColor(0x1A, 0x1A, 0x2E)  # bleu nuit
    COLOR_SUBTITLE    = RGBColor(0x6B, 0x6B, 0x6B)  # gris chaud
    COLOR_META        = RGBColor(0x99, 0x99, 0x99)  # gris clair
    COLOR_BODY        = RGBColor(0x2D, 0x2D, 0x2D)  # charbon
    COLOR_INTERVIEWER = RGBColor(0x2B, 0x4C, 0x6F)  # bleu acier
    COLOR_INVITE      = RGBColor(0x6B, 0x2D, 0x3E)  # bordeaux
    COLOR_FOOTER      = RGBColor(0xAA, 0xAA, 0xAA)  # gris léger
    RULE_COLOR        = "B8A88A"                     # or mat (hex)
    SEPARATOR_COLOR   = "D5D0C8"                     # beige gris (hex)
    FONT_MAIN         = "Georgia"

    doc = Document()

    # ── Métadonnées du document ──────────────────────────────────────────
    doc.core_properties.title = f"Entretien avec {', '.join(invites)}"
    doc.core_properties.author = ", ".join(interviewers)
    doc.core_properties.subject = "Transcription d'interview"
    doc.core_properties.category = "Interview"

    # ── Configuration de la page ─────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True

    # ── En-tête sur les pages de contenu (pas la page de garde) ──────────
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"Entretien avec {', '.join(invites)}")
    hr.font.size = Pt(8)
    hr.font.color.rgb = COLOR_META
    hr.font.name = FONT_MAIN
    hr.italic = True
    # Filet fin sous l'en-tête
    hpPr = hp._element.get_or_add_pPr()
    hpBdr = OxmlElement("w:pBdr")
    hbot = OxmlElement("w:bottom")
    hbot.set(qn("w:val"), "single")
    hbot.set(qn("w:sz"), "4")
    hbot.set(qn("w:space"), "4")
    hbot.set(qn("w:color"), SEPARATOR_COLOR)
    hpBdr.append(hbot)
    hpPr.append(hpBdr)

    # ── Pied de page avec numéro (pages de contenu) ──────────────────────
    footer_section = section.footer
    footer_section.is_linked_to_previous = False
    fp = footer_section.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Filet fin au-dessus du pied de page
    fpPr = fp._element.get_or_add_pPr()
    fpBdr = OxmlElement("w:pBdr")
    ftop = OxmlElement("w:top")
    ftop.set(qn("w:val"), "single")
    ftop.set(qn("w:sz"), "4")
    ftop.set(qn("w:space"), "6")
    ftop.set(qn("w:color"), SEPARATOR_COLOR)
    fpBdr.append(ftop)
    fpPr.append(fpBdr)
    # « — N — »
    r_dash1 = fp.add_run("— ")
    r_dash1.font.size = Pt(8)
    r_dash1.font.color.rgb = COLOR_FOOTER
    r_dash1.font.name = FONT_MAIN
    page_run = fp.add_run()
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = COLOR_FOOTER
    page_run.font.name = FONT_MAIN
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    page_run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    page_run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_end)
    r_dash2 = fp.add_run(" —")
    r_dash2.font.size = Pt(8)
    r_dash2.font.color.rgb = COLOR_FOOTER
    r_dash2.font.name = FONT_MAIN

    # ── Style Normal ─────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_MAIN
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BODY
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.35
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Helpers ──────────────────────────────────────────────────────────

    def add_rule(indent_cm=4.0, thickness=6, color=RULE_COLOR,
                 space_before=8, space_after=8):
        """Ajoute un filet horizontal décoratif via bordure de paragraphe."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.right_indent = Cm(indent_cm)
        r = p.add_run("\u00A0")
        r.font.size = Pt(2)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(thickness))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_spacer(height_pt=18):
        """Ajoute un espace vertical invisible."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run("\u00A0")
        r.font.size = Pt(height_pt)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def set_letter_spacing(run, pt_value):
        """Ajoute un espacement entre les caractères (en points)."""
        twips = int(pt_value * 20)
        rPr = run._r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run._r.insert(0, rPr)
        existing = rPr.find(qn("w:spacing"))
        if existing is not None:
            rPr.remove(existing)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(twips))
        rPr.append(spacing)

    # Date en français
    months_fr = [
        "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre",
    ]
    today = datetime.date.today()
    date_str = f"{today.day} {months_fr[today.month - 1]} {today.year}"

    # ═══════════════════════════════════════════════════════════════════
    #  PAGE DE GARDE
    # ═══════════════════════════════════════════════════════════════════

    if logo_path and os.path.exists(logo_path):
        # Avec logo : moins d'espace vertical, logo en haut
        for _ in range(3):
            add_spacer(20)

        # Logo centré
        doc.add_picture(logo_path, width=Cm(7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

        if is_antithese:
            # Sous-titre « Bon pour la tête »
            chan_para = doc.add_paragraph()
            chan_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chan_para.paragraph_format.space_before = Pt(2)
            chan_para.paragraph_format.space_after = Pt(10)
            chan_run = chan_para.add_run("Bon pour la tête")
            chan_run.font.size = Pt(10)
            chan_run.font.color.rgb = COLOR_SUBTITLE
            chan_run.font.name = FONT_MAIN
            chan_run.italic = True
    else:
        # Sans logo : espace vertical classique
        for _ in range(5):
            add_spacer(24)

    # Filet décoratif supérieur
    add_rule(indent_cm=4.0, thickness=6, space_before=0, space_after=12)

    # Titre de l'épisode (si fourni, typiquement en mode playlist)
    if titre:
        titre_para = doc.add_paragraph()
        titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titre_para.paragraph_format.space_before = Pt(6)
        titre_para.paragraph_format.space_after = Pt(8)
        titre_run = titre_para.add_run(titre)
        titre_run.font.size = Pt(16)
        titre_run.font.color.rgb = COLOR_TITLE
        titre_run.font.name = FONT_MAIN
        titre_run.italic = True

    # « Entretien avec » en petites capitales espacées
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.space_before = Pt(8)
    label_para.paragraph_format.space_after = Pt(4)
    label_run = label_para.add_run("Entretien avec")
    label_run.font.size = Pt(13)
    label_run.font.color.rgb = COLOR_SUBTITLE
    label_run.font.name = FONT_MAIN
    label_run.font.small_caps = True
    set_letter_spacing(label_run, 3)

    # Nom(s) de(s) invité(s) — grand titre
    invites_display = " & ".join(invites) if len(invites) > 1 else invites[0]
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(4)
    name_para.paragraph_format.space_after = Pt(12)
    name_run = name_para.add_run(invites_display)
    name_run.bold = True
    name_run.font.size = Pt(28)
    name_run.font.color.rgb = COLOR_TITLE
    name_run.font.name = FONT_MAIN

    # Filet décoratif inférieur
    add_rule(indent_cm=4.0, thickness=6, space_before=0, space_after=20)

    # Interviewers
    interviewers_display = " et ".join(interviewers)
    int_para = doc.add_paragraph()
    int_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    int_para.paragraph_format.space_before = Pt(8)
    int_para.paragraph_format.space_after = Pt(6)
    int_run = int_para.add_run(f"Interview menée par {interviewers_display}")
    int_run.font.size = Pt(12)
    int_run.font.color.rgb = COLOR_SUBTITLE
    int_run.font.name = FONT_MAIN
    int_run.italic = True

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(12)
    date_para.paragraph_format.space_after = Pt(2)
    date_run = date_para.add_run(f"Transcription du {date_str}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = COLOR_META
    date_run.font.name = FONT_MAIN

    # Source URL
    if video_url:
        url_para = doc.add_paragraph()
        url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        url_para.paragraph_format.space_before = Pt(2)
        url_para.paragraph_format.space_after = Pt(0)
        url_run = url_para.add_run(video_url)
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = COLOR_META
        url_run.font.name = FONT_MAIN
        url_run.italic = True

    # Saut de page → contenu
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    #  CONTENU DE L'INTERVIEW
    # ═══════════════════════════════════════════════════════════════════

    name_pattern = re.compile(r"^\*\*(.+?)\*\*\s*:?\s*(.*)")

    def flush_speaker(name, paragraphs, is_first_turn):
        """Écrit les paragraphes d'un locuteur dans le document."""
        for idx, text in enumerate(paragraphs):
            if not text.strip():
                continue

            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.35
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if idx == 0:
                # Premier paragraphe : nom en gras + texte
                para.paragraph_format.space_before = Pt(4 if is_first_turn else 14)
                para.paragraph_format.space_after = Pt(3)

                # Nom avec espace insécable avant le deux-points
                name_run = para.add_run(f"{name}\u00A0: ")
                name_run.bold = True
                name_run.font.size = Pt(11)
                name_run.font.name = FONT_MAIN
                if name in interviewers:
                    name_run.font.color.rgb = COLOR_INTERVIEWER
                else:
                    name_run.font.color.rgb = COLOR_INVITE

                text_run = para.add_run(text.strip())
                text_run.font.size = Pt(11)
                text_run.font.name = FONT_MAIN
                text_run.font.color.rgb = COLOR_BODY
            else:
                # Paragraphes de continuation (même locuteur, sans nom)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(3)

                text_run = para.add_run(text.strip())
                text_run.font.size = Pt(11)
                text_run.font.name = FONT_MAIN
                text_run.font.color.rgb = COLOR_BODY

    # Parser le texte réécrit en conservant les sauts de paragraphe
    current_name = None
    current_paragraphs = []   # liste de paragraphes pour le locuteur courant
    current_lines = []        # lignes du paragraphe courant
    turn_count = 0

    for line in rewritten_text.split("\n"):
        match = name_pattern.match(line)
        if match:
            # Clore le paragraphe en cours
            if current_lines:
                current_paragraphs.append(" ".join(current_lines))
                current_lines = []
            # Écrire le locuteur précédent
            if current_name and current_paragraphs:
                flush_speaker(current_name, current_paragraphs, turn_count == 0)
                turn_count += 1
                current_paragraphs = []

            current_name = match.group(1)
            if match.group(2):
                current_lines = [match.group(2)]
        elif line.strip() == "":
            # Ligne vide = saut de paragraphe au sein de la même prise de parole
            if current_lines:
                current_paragraphs.append(" ".join(current_lines))
                current_lines = []
        else:
            current_lines.append(line.strip())

    # Flush final
    if current_lines:
        current_paragraphs.append(" ".join(current_lines))
    if current_name and current_paragraphs:
        flush_speaker(current_name, current_paragraphs, turn_count == 0)

    # ═══════════════════════════════════════════════════════════════════
    #  CRÉDIT EN FIN DE DOCUMENT
    # ═══════════════════════════════════════════════════════════════════

    add_spacer(12)
    add_rule(indent_cm=3.5, thickness=4, space_before=4, space_after=10)

    credit_para = doc.add_paragraph()
    credit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    credit_run = credit_para.add_run(
        "Transcription automatique (WhisperX + Pyannote)"
        " — Réécriture par Claude (Anthropic)"
    )
    credit_run.font.size = Pt(8)
    credit_run.font.color.rgb = COLOR_FOOTER
    credit_run.font.name = FONT_MAIN
    credit_run.italic = True

    doc.save(output_path)
    print(f"📄 Document sauvegardé : {output_path}")


# ---------------------------------------------------------------------------
# 11. Sauvegarde texte
# ---------------------------------------------------------------------------

def save_text(text: str, path: str, label: str):
    """Sauvegarde du texte dans un fichier."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"💾 {label} : {path}")


# ---------------------------------------------------------------------------
# 12. Logo Antithèse (téléchargement + conversion SVG → PNG)
# ---------------------------------------------------------------------------

def download_antithese_logo(workdir: str) -> str | None:
    """
    Télécharge le logo Antithèse/Bon pour la tête, le convertit en PNG
    pour intégration dans le docx. Utilise un cache dans le workdir.
    """
    cached = os.path.join(workdir, "logo_antithese.png")
    if os.path.exists(cached):
        return cached

    print("🎨 Récupération du logo Antithèse...")
    try:
        req = urllib.request.Request(
            ANTITHESE_LOGO_URL,
            headers={"User-Agent": "Mozilla/5.0"},
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            svg_text = resp.read().decode("utf-8", errors="replace")

        # Adapter pour fond clair : remplacer le blanc par du quasi-noir
        for pattern, repl in [
            (r'fill\s*:\s*#(?:fff(?:fff)?)\b', 'fill:#1a1a1a'),
            (r'fill\s*=\s*"#(?:fff(?:fff)?)"', 'fill="#1a1a1a"'),
            (r'fill\s*:\s*white\b', 'fill:#1a1a1a'),
            (r'fill\s*=\s*"white"', 'fill="#1a1a1a"'),
            (r'fill\s*:\s*rgb\(\s*255\s*,\s*255\s*,\s*255\s*\)', 'fill:#1a1a1a'),
            (r'fill\s*=\s*"rgb\(\s*255\s*,\s*255\s*,\s*255\s*\)"', 'fill="#1a1a1a"'),
        ]:
            svg_text = re.sub(pattern, repl, svg_text, flags=re.IGNORECASE)

        # Convertir SVG → PNG avec cairosvg
        try:
            import cairosvg
        except ImportError:
            print("   ⚠️  cairosvg non installé — logo non disponible")
            return None

        cairosvg.svg2png(
            bytestring=svg_text.encode("utf-8"),
            write_to=cached,
            output_width=800,
        )
        print("   ✅ Logo récupéré et converti en PNG")
        return cached

    except Exception as e:
        print(f"   ⚠️  Logo non disponible : {e}")
        return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

RESUME_STEPS = ["transcribe", "name", "heuristic", "claude-fix", "rewrite", "docx"]


def process_single_video(
    source: str,
    interviewers: list[str],
    invites: list[str],
    output_path: str,
    args,
    api_key: str,
    hf_token: str,
    is_antithese: bool = False,
    titre: str = None,
):
    """Exécute le pipeline complet pour une seule vidéo. Retourne (timer, api_stats)."""
    ignore_duration = args.ignore or args.skip
    start_from = RESUME_STEPS.index(args.resume) if args.resume else 0

    all_names = interviewers + invites
    num_speakers = len(all_names)

    workdir = get_workdir(output_path)

    print(f"\n{'='*60}")
    if is_antithese:
        print("🎬 Antithèse · Bon pour la tête — Transcription d'interview")
    else:
        print("🎬 Pipeline de transcription d'interview")
    print(f"{'='*60}")
    if source:
        print(f"   Source        : {source}")
    print(f"   Interviewers  : {', '.join(interviewers)}")
    print(f"   Invité(s)     : {', '.join(invites)}")
    print(f"   Total         : {num_speakers} locuteurs")
    print(f"   Modèle Claude : {args.model}")
    print(f"   Modèle Whisper: {args.whisper_model}")
    if ignore_duration:
        print(f"   Ignorer début : {ignore_duration}")
    if args.resume:
        print(f"   Reprise depuis: {args.resume}")
    print(f"   Workdir       : {workdir}")
    print(f"{'='*60}\n")

    timer = StepTimer()
    api_stats = APIStats()
    tmp_dir = None

    try:
        # ── Étape 1 : Transcription ──────────────────────────────────────
        if start_from <= RESUME_STEPS.index("transcribe"):
            timer.start("Téléchargement audio")
            tmp_dir = tempfile.mkdtemp(prefix="interview_")

            audio_path = download_audio(source, tmp_dir)

            if ignore_duration:
                audio_path = trim_audio(audio_path, ignore_duration, tmp_dir)

            timer.start("Transcription + Diarisation")
            segments = transcribe_and_diarize(
                audio_path, num_speakers, hf_token,
                whisper_model=args.whisper_model,
                batch_size=args.batch_size,
            )
            save_checkpoint(workdir, "transcription", segments, 1)
        else:
            print("📂 Chargement du checkpoint transcription...")
            segments = load_checkpoint(workdir, "transcription", 1)
            if segments is None:
                print("❌ Checkpoint 01_transcription.json introuvable dans le workdir")
                sys.exit(1)
            print(f"   ✅ {len(segments)} segments chargés")

        # ── Étape 2 : Attribution intelligente des noms ───────────────────
        if start_from <= RESUME_STEPS.index("name"):
            timer.start("Attribution des noms")
            segments = replace_speaker_names(segments, interviewers, invites)
            save_checkpoint(workdir, "named", segments, 2)

            # Sauvegarde transcription brute (avant corrections)
            raw_text = format_transcript_for_claude(segments)
            save_text(
                raw_text,
                str(Path(output_path).with_suffix(".1_brut.txt")),
                "Transcription brute",
            )
        else:
            print("📂 Chargement du checkpoint named...")
            segments = load_checkpoint(workdir, "named", 2)
            if segments is None:
                print("❌ Checkpoint 02_named.json introuvable dans le workdir")
                sys.exit(1)
            print(f"   ✅ {len(segments)} segments chargés")

        # Calculer les stats après attribution des noms
        speaker_stats = compute_speaker_stats(segments)
        speaker_stats_str = format_speaker_stats(speaker_stats)

        # ── Étape 3 : Corrections heuristiques ───────────────────────────
        if start_from <= RESUME_STEPS.index("heuristic"):
            timer.start("Corrections heuristiques")
            seg_count_before = len(segments)

            if not args.skip_heuristics:
                segments = fix_attributions_heuristic(segments, interviewers, invites)
            else:
                print("🔧 Corrections heuristiques : ignorées (--skip-heuristics)")

            save_checkpoint(workdir, "heuristic", segments, 3)

            heuristic_text = format_transcript_for_claude(segments)
            save_text(
                heuristic_text,
                str(Path(output_path).with_suffix(".2_corrige_heuristique.txt")),
                "Après corrections heuristiques",
            )
        else:
            print("📂 Chargement du checkpoint heuristic...")
            segments = load_checkpoint(workdir, "heuristic", 3)
            if segments is None:
                print("❌ Checkpoint 03_heuristic.json introuvable dans le workdir")
                sys.exit(1)
            heuristic_text = format_transcript_for_claude(segments)
            print(f"   ✅ {len(segments)} segments chargés")

        if args.raw_only:
            timer.stop()
            print(f"\n✅ Terminé (mode brut).")
            print(f"\n📊 Statistiques de parole :")
            print(speaker_stats_str)
            print(f"\n⏱️  {timer.summary()}")
            return timer, api_stats

        # ── Étape 4 : Correction des attributions par Claude ─────────────
        if start_from <= RESUME_STEPS.index("claude-fix"):
            if not args.skip_claude_fix:
                timer.start("Correction attributions (Claude)")
                corrected_text = fix_attributions_claude(
                    heuristic_text, interviewers, invites, api_key,
                    args.model, api_stats, speaker_stats_str,
                )
                save_text(
                    corrected_text,
                    str(Path(output_path).with_suffix(".3_corrige_claude.txt")),
                    "Après correction Claude",
                )
                save_checkpoint(workdir, "claude_fix", {"text": corrected_text}, 4)
            else:
                corrected_text = heuristic_text
                print("🔍 Correction Claude : ignorée (--skip-claude-fix)")
                save_checkpoint(workdir, "claude_fix", {"text": corrected_text}, 4)
        else:
            print("📂 Chargement du checkpoint claude_fix...")
            cp = load_checkpoint(workdir, "claude_fix", 4)
            if cp is None:
                print("❌ Checkpoint 04_claude_fix.json introuvable dans le workdir")
                sys.exit(1)
            corrected_text = cp["text"]
            print(f"   ✅ Texte corrigé chargé ({len(corrected_text)} caractères)")

        # ── Étape 5 : Réécriture avec Claude ─────────────────────────────
        if start_from <= RESUME_STEPS.index("rewrite"):
            timer.start("Réécriture (Claude)")
            rewritten = rewrite_with_claude(
                corrected_text, interviewers, invites, api_key,
                args.model, api_stats,
            )
            save_text(
                rewritten,
                str(Path(output_path).with_suffix(".4_reecrit.txt")),
                "Texte réécrit",
            )
            save_checkpoint(workdir, "rewrite", {"text": rewritten}, 5)
        else:
            print("📂 Chargement du checkpoint rewrite...")
            cp = load_checkpoint(workdir, "rewrite", 5)
            if cp is None:
                print("❌ Checkpoint 05_rewrite.json introuvable dans le workdir")
                sys.exit(1)
            rewritten = cp["text"]
            print(f"   ✅ Texte réécrit chargé ({len(rewritten)} caractères)")

        # ── Étape 6 : Génération du .docx ────────────────────────────────
        timer.start("Génération DOCX")
        video_url = source if source and source.startswith("http") else None

        # Logo Antithèse si interviewers par défaut
        logo_path = None
        if is_antithese:
            logo_path = download_antithese_logo(workdir)

        generate_docx(
            rewritten, interviewers, invites, output_path, video_url,
            logo_path=logo_path, is_antithese=is_antithese, titre=titre,
        )

    finally:
        # Cleanup des fichiers temporaires audio
        if tmp_dir and os.path.exists(tmp_dir):
            import shutil
            try:
                shutil.rmtree(tmp_dir)
            except OSError:
                pass

    timer.stop()

    # ── Résumé final ─────────────────────────────────────────────────────
    print(f"\n{'='*60}")
    print("🎉 Terminé !")
    print(f"{'='*60}")
    print(f"   📄 Document final                : {output_path}")
    print(f"   📝 Texte réécrit                 : {Path(output_path).with_suffix('.4_reecrit.txt')}")
    print(f"   🔍 Après correction Claude       : {Path(output_path).with_suffix('.3_corrige_claude.txt')}")
    print(f"   🔧 Après correction heuristique  : {Path(output_path).with_suffix('.2_corrige_heuristique.txt')}")
    print(f"   📋 Transcription brute           : {Path(output_path).with_suffix('.1_brut.txt')}")
    print(f"   💾 Workdir (checkpoints)         : {workdir}")

    print(f"\n📊 Statistiques de parole :")
    if segments:
        final_stats = compute_speaker_stats(segments)
        print(format_speaker_stats(final_stats))

    if api_stats.calls > 0:
        print(f"\n🤖 API Claude :")
        print(api_stats.summary())

    print(f"\n⏱️  Durée par étape :")
    print(timer.summary())
    print(f"{'='*60}\n")

    return timer, api_stats


def main():
    parser = argparse.ArgumentParser(
        description="Transcription d'interviews YouTube avec diarisation et réécriture.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples :
  # Interview Antithèse (interviewers par défaut : Martin Bernard, Ouriel Barbezat)
  python transcrire.py "https://youtube.com/watch?v=xxx" --invites "Pierre Gallaz"

  # Avec bande-annonce de 2min30 à ignorer
  python transcrire.py "https://youtube.com/watch?v=xxx" --invites "Sophie" --skip 2:30

  # Interviewers personnalisés (pas de logo Antithèse)
  python transcrire.py interview.mp4 \\
      --hosts "Jean-Marc" --invites "Sophie" "Paul" --output entretien.docx

  # Transcription brute seulement
  python transcrire.py interview.mp3 --invites "Sophie" --raw-only

  # Reprendre depuis la réécriture
  python transcrire.py source --invites "Sophie" --output entretien.docx --resume rewrite

  # Mode playlist : sélection interactive depuis la playlist Antithèse
  python transcrire.py --playlist

Variables d'environnement requises :
  ANTHROPIC_API_KEY : Clé API Anthropic (pour Claude)
  HF_TOKEN          : Token Hugging Face (pour Pyannote)
        """,
    )

    parser.add_argument(
        "source", nargs="?", default=None,
        help="URL YouTube ou chemin vers un fichier audio local (optionnel si --resume)",
    )
    parser.add_argument(
        "--interviewers", "--hosts", "-i", nargs="+", default=None,
        help="Nom(s) des interviewers (défaut: Martin Bernard, Ouriel Barbezat)",
    )
    parser.add_argument(
        "--invites", "-g", nargs="+", default=None,
        help="Nom(s) des invité(s) (requis sauf en mode --playlist)",
    )
    parser.add_argument("--output", "-o", default=None, help="Chemin du fichier .docx de sortie")
    parser.add_argument(
        "--playlist", action="store_true",
        help="Mode playlist : lit la playlist Antithèse, sélection interactive, batch automatique",
    )

    # Ignorer le début
    parser.add_argument(
        "--ignore", default=None, metavar="MM:SS",
        help="Ignorer les N premières minutes:secondes (ex: 2:30)",
    )
    parser.add_argument(
        "--skip", default=None, metavar="MM:SS",
        help="Alias pour --ignore",
    )

    # Modes partiels
    parser.add_argument("--raw-only", action="store_true",
                        help="Seulement transcrire + corriger les attributions, pas de réécriture")
    parser.add_argument("--skip-heuristics", action="store_true",
                        help="Ne pas appliquer les corrections heuristiques")
    parser.add_argument("--skip-claude-fix", action="store_true",
                        help="Ne pas faire la passe de correction d'attributions par Claude")

    # Reprise
    parser.add_argument(
        "--resume", default=None, choices=RESUME_STEPS,
        help="Reprendre depuis une étape (name, heuristic, claude-fix, rewrite, docx)",
    )
    parser.add_argument(
        "--rewrite-only", action="store_true",
        help="Reprendre depuis l'étape de réécriture (équivalent à --resume rewrite)",
    )

    # Modèles
    parser.add_argument("--model", default="claude-opus-4-5",
                        help="Modèle Claude à utiliser (défaut: claude-opus-4-5, "
                             "cf. A/B 2026-05-25 dans traduire.py)")
    parser.add_argument("--whisper-model", default="large-v3",
                        help="Modèle Whisper (tiny, base, small, medium, large-v3)")
    parser.add_argument("--batch-size", type=int, default=16,
                        help="Taille du batch Whisper (réduire si CUDA OOM)")

    args = parser.parse_args()

    # --rewrite-only = --resume rewrite
    if args.rewrite_only:
        args.resume = "rewrite"

    # Validation : --invites requis sauf en mode --playlist
    if not args.playlist and not args.invites:
        parser.error("l'argument --invites/-g est requis (sauf en mode --playlist)")

    # Vérification des clés API
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    hf_token = os.environ.get("HF_TOKEN")

    # ── Mode playlist ────────────────────────────────────────────────────
    if args.playlist:
        if args.resume:
            print("⚠️  --resume est ignoré en mode --playlist")
            args.resume = None

        videos = fetch_playlist()
        display_playlist(videos)
        selected = select_videos(videos)

        print(f"\n🎬 {len(selected)} vidéo(s) sélectionnée(s)\n")

        interviewers = list(DEFAULT_INTERVIEWERS)

        if not hf_token:
            print("❌ Variable d'environnement HF_TOKEN manquante.")
            print("   Créez un token sur https://huggingface.co/settings/tokens")
            sys.exit(1)

        if not api_key and not args.raw_only:
            print("❌ Variable d'environnement ANTHROPIC_API_KEY manquante.")
            sys.exit(1)

        for i, video in enumerate(selected, 1):
            print(f"\n{'='*60}")
            print(f"📹 Vidéo {i}/{len(selected)} : {video['title']}")
            print(f"   Invité(s) : {', '.join(video['guests'])}")
            print(f"{'='*60}")

            guest_slug = "_".join(video["guests"]).replace(" ", "-")
            output_path = f"interview_{guest_slug}.docx"

            process_single_video(
                source=video["url"],
                interviewers=interviewers,
                invites=video["guests"],
                output_path=output_path,
                args=args,
                api_key=api_key,
                hf_token=hf_token,
                is_antithese=True,
                titre=video["title"],
            )

        print(f"\n{'='*60}")
        print(f"🎉 Playlist terminée — {len(selected)} vidéo(s) traitée(s)")
        print(f"{'='*60}\n")
        return

    # ── Mode vidéo unique ────────────────────────────────────────────────
    # --skip = alias pour --ignore
    ignore_duration = args.ignore or args.skip

    # Déterminer l'étape de départ
    start_from = RESUME_STEPS.index(args.resume) if args.resume else 0
    needs_audio = start_from == 0

    # Vérifier que source est fourni si on a besoin de l'audio
    if needs_audio and not args.source:
        print("❌ Source requise (URL YouTube ou fichier audio) sauf avec --resume")
        sys.exit(1)

    if needs_audio and not hf_token:
        print("❌ Variable d'environnement HF_TOKEN manquante.")
        print("   Créez un token sur https://huggingface.co/settings/tokens")
        sys.exit(1)

    needs_claude = not args.raw_only and start_from < RESUME_STEPS.index("docx")
    if needs_claude and not api_key:
        print("❌ Variable d'environnement ANTHROPIC_API_KEY manquante.")
        sys.exit(1)

    # Interviewers : défaut Antithèse si non spécifié
    if args.interviewers is None:
        interviewers = list(DEFAULT_INTERVIEWERS)
        is_antithese = True
    else:
        interviewers = args.interviewers
        is_antithese = False

    invites = args.invites

    if args.output is None:
        guest_slug = "_".join(invites).replace(" ", "-")
        args.output = f"interview_{guest_slug}.docx"

    process_single_video(
        source=args.source,
        interviewers=interviewers,
        invites=invites,
        output_path=args.output,
        args=args,
        api_key=api_key,
        hf_token=hf_token,
        is_antithese=is_antithese,
    )


if __name__ == "__main__":
    main()
